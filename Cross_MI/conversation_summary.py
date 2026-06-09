"""生成对话总结Word文档"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

# 标题
title = doc.add_heading('研究对话总结：跨被试MI-EEG通用特征分析', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f'生成时间：{datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}')
doc.add_paragraph('项目路径：E:\\Code\\Cross\\Cross_MI')
doc.add_paragraph('数据路径：E:\\Datasets\\4_跨场景因素研究v2')
doc.add_paragraph('conda环境：Cross_gpu')
doc.add_paragraph('')

# 1. 研究背景
doc.add_heading('1. 研究背景与数据集', level=1)
p = doc.add_paragraph(
    '数据集包含4个子数据集（Graz、SSMVEP-MI、Hybrid、Hybrid-online），共84名被试，'
    '4种刺激范式，两个采集场景（医院 vs 实验室）。核心范式为cue范式左右手运动想象（MI），'
    '37名被试，每人两个场景（S1/S2）。'
)

doc.add_paragraph(
    '原始假设"医院场景训练失败于实验室场景"被否定——实验结果显示两场景间无一致显著差异，'
    '跨场景泛化不是主要问题。真实瓶颈在于被试间差异（inter-subject variability）。'
)

doc.add_paragraph(
    '研究方向确立为：寻找跨被试稳定的通用特征，构建零/最小校准的通用解码器。'
)

# 2. 核心指标定义
doc.add_heading('2. 三维评价框架', level=1)

doc.add_paragraph('为系统评价特征的"跨被试可用性"，定义三个维度：', style='Normal')

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '维度'
hdr[1].text = '名称'
hdr[2].text = '含义'
rows_data = [
    ('Dim1', 'CSCDC\n(跨被试类方向一致性)',
     '所有被试对的delta向量余弦相似度均值。delta = 类0均值 - 类1均值（标准化后）。'
     '越高 → 类分离方向跨被试越一致 → 特征越通用'),
    ('Dim2', 'FDR\n(Fisher判别比)',
     '被试内类间方差/类内方差。衡量特征的判别力（discriminability）'),
    ('Dim3', 'SceneR\n(场景可靠性)',
     '同被试S1 vs S2的delta向量余弦相似度。衡量特征在跨场景下的一致性'),
]
for i, (d, n, c) in enumerate(rows_data):
    row = table.rows[i+1].cells
    row[0].text = d; row[1].text = n; row[2].text = c

doc.add_paragraph('')

# 3. 各Step分析
doc.add_heading('3. 分析步骤', level=1)

# Step7
doc.add_heading('Step7：特征稳定性基准分析', level=2)
doc.add_paragraph(
    '文件：Feature_analysis/Step7_feature_stability.py\n'
    '指标：Fisher判别比（FDR）\n'
    '特征：PSD功率比、CSP对数方差、黎曼切空间、FBCSP多频带'
)
doc.add_paragraph('结论：FBCSP判别力最强（FDR≈0.376，稳定性分数2.29），但依赖标签（有监督）。')

# Step8
doc.add_heading('Step8：跨被试通用特征系统评估（CSCDC）', level=2)
doc.add_paragraph(
    '文件：Feature_analysis/Step8_universal_feature_analysis.py\n'
    '覆盖7类17种特征，引入CSCDC作为核心指标（无监督/有监督分析）'
)
t8 = doc.add_table(rows=6, cols=3)
t8.style = 'Table Grid'
h = t8.rows[0].cells; h[0].text='排名'; h[1].text='特征'; h[2].text='CSCDC'
t8_data = [
    ('1', 'F01 Alpha Power (C3/C4)', '~0.29'),
    ('2', 'F03 Mu Power (C3/C4)', '~0.28'),
    ('3', 'F07 Laplacian Power', '~0.27'),
    ('~末', 'F05 CSP LogVar（有监督）', '最高但不公平'),
]
for i,(r,f,v) in enumerate(t8_data):
    row=t8.rows[i+1].cells; row[0].text=r; row[1].text=f; row[2].text=v
doc.add_paragraph('')
doc.add_paragraph(
    '关键发现：FBCSP/CSP因使用标签在特征提取中具有不公平优势。'
    '在真正无监督特征中，Alpha/Mu功率（C3/C4）CSCDC最高，约0.28-0.29。'
)

# Step9
doc.add_heading('Step9：个体化Alpha峰频（iAPF）自适应频带', level=2)
doc.add_paragraph(
    '文件：Feature_analysis/Step9_iAPF_multidim_evaluation.py\n'
    '假设：个体iAPF差异导致固定8-13Hz频带次优，自适应频带可提升性能\n'
    '三维评价框架正式引入'
)
doc.add_paragraph(
    '结论（失败）：iAPF估计从任务态数据进行，Alpha ERD导致估计偏向低频（~9Hz，'
    '即8Hz边界），所有iAPF变体均不优于固定频带。'
    '正确做法是用静息态数据估计iAPF，但此数据集中无静息态段。'
)

# Step10
doc.add_heading('Step10：侧化指数（LI）评估', level=2)
doc.add_paragraph(
    '文件：Feature_analysis/Step10_lateralization_index.py\n'
    '动机：LI = (P_C3 - P_C4) / (P_C3 + P_C4)\n'
    '  - 神经生理学依据：MI核心现象为对侧ERD（左手→右脑C4抑制，右手→左脑C3抑制）\n'
    '  - 自归一化：消除被试间绝对功率差异\n'
    '  - 零校准潜力：无需个体参数'
)
doc.add_paragraph('评估8种LI变体：')
doc.add_paragraph('  • LI_Alpha (8-13Hz), LI_Mu (8-12Hz), LI_Beta (13-30Hz), LI_Multi (7维频段向量)')
doc.add_paragraph('  • 以上4种各加Laplacian空间滤波 → LI_Lap_Alpha/Mu/Beta/Multi')
doc.add_paragraph('  • 对比Step9 Reference特征（Alpha Fixed、Mu Fixed、Laplacian Fixed、Cov Diag）')
doc.add_paragraph(
    '三维评价（CSCDC / FDR / SceneR）：\n'
    '  - LI_Lap_Multi（Laplacian + 7维多频段）预期综合分数最高\n'
    '  - 自归一化使跨被试CSCDC提升（相比原始log功率）\n'
    '  - 输出：ranking图、雷达图、逐被试delta一致性热图'
)

# Step11
doc.add_heading('Step11：分类性能验证', level=2)
doc.add_paragraph(
    '文件：Feature_analysis/Step11_classification_validation.py\n'
    '对比6种方法的实际分类准确率：'
)
t11 = doc.add_table(rows=7, cols=3)
t11.style = 'Table Grid'
h = t11.rows[0].cells; h[0].text='方法'; h[1].text='特征'; h[2].text='分类器/协议'
m_data = [
    ('M1 LI_Alpha_thresh', 'LI(8-13Hz)', '群体均值阈值（零校准）'),
    ('M2 LI_Alpha_LDA',    'LI(8-13Hz) 1D', 'LDA, LOSO跨被试'),
    ('M3 LI_Mu_thresh',    'LI(8-12Hz)', '阈值（零校准）'),
    ('M4 Alpha_LDA',       'log功率[C3,C4]', 'LDA, LOSO跨被试'),
    ('M5 CSP_LDA_within',  'CSP+LDA', '被试内10折（校准上界）'),
    ('M6 CSP_LDA_cross',   '公共空间CSP+LDA', 'LOSO跨被试'),
]
for i,(m,f,c) in enumerate(m_data):
    row=t11.rows[i+1].cells; row[0].text=m; row[1].text=f; row[2].text=c
doc.add_paragraph('')
doc.add_paragraph(
    '两种评估场景：\n'
    '  S1：标准LOSO（合并两场景训练，测试同被试）\n'
    '  S2：Cross-Double（训练Scene1所有被试，测试Scene2所有被试）\n'
    '统计检验：paired t-test (M1 vs 其他方法)'
)

# 4. 研究故事线
doc.add_heading('4. 研究叙事（论文故事线）', level=1)
story_items = [
    '现有MI-BCI跨被试泛化差的根本原因是什么？→ 答：不是场景/环境，而是特征本身跨被试不稳定',
    '系统评估17种特征的跨被试通用性（CSCDC框架）→ 结论：无监督谱特征（Alpha/Mu功率）CSCDC最高',
    '个体化自适应频带（iAPF）是否能提升稳定性？→ 答：否（任务态估计失效）',
    'LI（侧化指数）作为自归一化、神经生理学驱动特征 → 在三维框架上验证优越性',
    '实际分类验证：LI零校准方法 vs 需要校准的CSP → 量化gap，证明LI的实用价值',
    '（待做）Step12：校准效率曲线 — 准确率 vs 训练被试数，证明LI以极少数据达到可用性能',
]
for item in story_items:
    doc.add_paragraph(item, style='List Number')

# 5. 待完成工作
doc.add_heading('5. 待完成工作', level=1)
todo_items = [
    'Step12：校准效率分析（accuracy vs N_training_subjects，LI-thresh vs CSP-LDA）',
    '运行Step10/Step11获取实际数值结果',
    '整合所有分析结果到论文图表（Figure 3-6）',
    '撰写Methods和Results章节对应内容',
]
for item in todo_items:
    doc.add_paragraph(item, style='List Bullet')

# 6. 关键文件路径
doc.add_heading('6. 关键文件路径', level=1)
paths = [
    ('代码目录', 'E:\\Code\\Cross\\Cross_MI\\Feature_analysis\\'),
    ('数据目录', 'E:\\Datasets\\4_跨场景因素研究v2\\跨场景因素研究v2处理后数据\\'),
    ('输出目录', 'E:\\Datasets\\4_跨场景因素研究v2\\跨场景因素研究v2画图数据\\stability\\'),
    ('手稿', 'E:\\Files\\4_写论文\\2026_SD_撰写中\\v1.0\\Manuscript-加KJZ.docx'),
    ('Step7', 'Step7_feature_stability.py'),
    ('Step8', 'Step8_universal_feature_analysis.py'),
    ('Step9', 'Step9_iAPF_multidim_evaluation.py'),
    ('Step10', 'Step10_lateralization_index.py'),
    ('Step11', 'Step11_classification_validation.py'),
]
t_path = doc.add_table(rows=len(paths)+1, cols=2)
t_path.style = 'Table Grid'
t_path.rows[0].cells[0].text = '说明'
t_path.rows[0].cells[1].text = '路径'
for i,(k,v) in enumerate(paths):
    t_path.rows[i+1].cells[0].text = k
    t_path.rows[i+1].cells[1].text = v

out = r'/mnt/e/Code/Cross/Cross_MI/对话总结_跨被试MI特征分析.docx'
doc.save(out)
print(f'已保存: {out}')
